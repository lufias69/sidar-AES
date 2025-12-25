"""
Convert MANUSCRIPT_DRAFT_v1.md to AJET-compliant Word document
Applies AJET formatting requirements:
- Calibri 10pt for body text
- Arial 14pt bold for title
- Arial 12pt bold for Level 1 headings
- Arial 10pt bold for Level 2 headings
- 3.0 cm margins all sides
- Single spacing, left justified
- Abstract: indented 1.0 cm, no heading
- Implications: indented 1.0 cm, italic heading
- References: hanging indent 0.5 cm
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys

def set_run_font(run, font_name, font_size, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # Ensure font is applied to complex scripts too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def set_paragraph_spacing(paragraph, line_spacing=1.0, space_after=0, space_before=0):
    """Set paragraph spacing"""
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)

def add_hanging_indent(paragraph, indent_cm=0.5):
    """Add hanging indent for references"""
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    paragraph.paragraph_format.first_line_indent = Cm(-indent_cm)

def add_left_indent(paragraph, indent_cm=1.0):
    """Add left indent (for abstract, implications)"""
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    paragraph.paragraph_format.right_indent = Cm(indent_cm)

def process_markdown_formatting(text):
    """Extract text with bold/italic markers"""
    # Return list of (text, is_bold, is_italic) tuples
    segments = []
    current_pos = 0
    
    # Find all **text** and *text* patterns
    bold_pattern = r'\*\*(.+?)\*\*'
    italic_pattern = r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)'
    
    # Process bold first
    for match in re.finditer(bold_pattern, text):
        # Add text before match
        if match.start() > current_pos:
            segments.append((text[current_pos:match.start()], False, False))
        # Add bold text
        segments.append((match.group(1), True, False))
        current_pos = match.end()
    
    # Add remaining text
    if current_pos < len(text):
        remaining = text[current_pos:]
        # Check for italics in remaining
        for match in re.finditer(italic_pattern, remaining):
            if match.start() > 0:
                segments.append((remaining[:match.start()], False, False))
            segments.append((match.group(1), False, True))
            remaining = remaining[match.end():]
        if remaining:
            segments.append((remaining, False, False))
    
    return segments if segments else [(text, False, False)]

def parse_inline_formatting(text):
    """Parse bold and italic markers, return list of (text, bold, italic) tuples"""
    segments = []
    i = 0
    while i < len(text):
        # Check for **bold**
        if text[i:i+2] == '**':
            end = text.find('**', i+2)
            if end != -1:
                segments.append((text[i+2:end], True, False))
                i = end + 2
                continue
        # Check for *italic*
        if text[i] == '*':
            end = text.find('*', i+1)
            if end != -1:
                segments.append((text[i+1:end], False, True))
                i = end + 1
                continue
        # Regular text - find next marker
        next_bold = text.find('**', i)
        next_italic = text.find('*', i)
        
        if next_bold == -1 and next_italic == -1:
            segments.append((text[i:], False, False))
            break
        elif next_bold == -1:
            segments.append((text[i:next_italic], False, False))
            i = next_italic
        elif next_italic == -1:
            segments.append((text[i:next_bold], False, False))
            i = next_bold
        else:
            next_marker = min(next_bold, next_italic)
            segments.append((text[i:next_marker], False, False))
            i = next_marker
    
    return segments if segments else [(text, False, False)]

def parse_markdown_table(lines, start_idx):
    """Parse markdown table starting at start_idx"""
    table_lines = []
    i = start_idx
    
    # Collect table lines
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1
    
    if len(table_lines) < 3:  # Need header, separator, at least one row
        return None, start_idx
    
    # Parse header
    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Parse rows (skip separator line at index 1)
    rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    
    return (header, rows), i

def add_table_to_doc(doc, header, rows):
    """Add table with AJET formatting"""
    table = doc.add_table(rows=1+len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    
    # Header row - bold
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        set_run_font(run, 'Calibri', 10, bold=True)
        p.paragraph_format.space_after = Pt(0)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            set_run_font(run, 'Calibri', 10)
            p.paragraph_format.space_after = Pt(0)
    
    return table

def convert_markdown_to_docx(md_file, output_file):
    """Convert markdown manuscript to AJET Word format"""
    
    # Read file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create document
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
    
    # State tracking
    in_abstract = False
    in_implications = False
    in_references = False
    skip_lines = 0
    
    i = 0
    while i < len(lines):
        if skip_lines > 0:
            skip_lines -= 1
            i += 1
            continue
            
        line = lines[i].rstrip()
        
        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue
        
        # Title (first # heading)
        if line.startswith('# ') and i < 5:
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            set_run_font(run, 'Arial', 14, bold=True)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.space_before = Pt(0)
            i += 1
            continue
        
        # Running title
        if line.startswith('**Running Title**'):
            # Skip running title - not needed in manuscript body
            i += 1
            continue
        
        # ABSTRACT section
        if line.strip() == '## ABSTRACT':
            in_abstract = True
            # Don't add heading - AJET requires no "Abstract" heading
            i += 1
            continue
        
        # IMPLICATIONS section  
        if line.strip() == '## IMPLICATIONS FOR PRACTICE OR POLICY':
            in_implications = True
            in_abstract = False
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            run = p.add_run('Implications for practice or policy')
            set_run_font(run, 'Calibri', 10, italic=True)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(12)
            i += 1
            continue
        
        # RESEARCH HIGHLIGHTS - skip this section
        if line.strip() == '## RESEARCH HIGHLIGHTS':
            # Skip until next ## heading
            i += 1
            while i < len(lines) and not (lines[i].startswith('##') and not lines[i].startswith('###')):
                i += 1
            continue
        
        # REFERENCES section
        if line.strip() == '## REFERENCES':
            in_references = True
            in_implications = False
            in_abstract = False
            p = doc.add_paragraph()
            run = p.add_run('REFERENCES')
            set_run_font(run, 'Arial', 12, bold=True)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(12)
            i += 1
            continue
        
        # Level 1 heading (##)
        if line.startswith('## '):
            in_abstract = False
            in_implications = False
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            set_run_font(run, 'Arial', 12, bold=True)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(12)
            i += 1
            continue
        
        # Level 2 heading (###)
        if line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            set_run_font(run, 'Arial', 10, bold=True)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(6)
            i += 1
            continue
        
        # Tables
        if line.strip().startswith('|'):
            table_data, end_idx = parse_markdown_table(lines, i)
            if table_data:
                header, rows = table_data
                add_table_to_doc(doc, header, rows)
                i = end_idx
                continue
        
        # Figure references (![...](...))
        if line.strip().startswith('!['):
            # Skip figure placeholders - user will insert manually
            i += 1
            continue
        
        # Bullet points in implications
        if in_implications and line.strip().startswith('- '):
            bullet_text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            
            # Parse inline formatting
            segments = parse_inline_formatting(bullet_text)
            for text, is_bold, is_italic in segments:
                run = p.add_run(text)
                set_run_font(run, 'Calibri', 10, bold=is_bold, italic=is_italic)
            
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        
        # Regular bullet points (skip RESEARCH HIGHLIGHTS which we already skipped)
        if line.strip().startswith('•'):
            # Skip
            i += 1
            continue
        
        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            
            # Apply indentation for abstract
            if in_abstract:
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.right_indent = Cm(1.0)
            
            # Apply hanging indent for references
            if in_references:
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.first_line_indent = Cm(-0.5)
            
            # Parse inline formatting
            segments = parse_inline_formatting(line.strip())
            for text, is_bold, is_italic in segments:
                run = p.add_run(text)
                set_run_font(run, 'Calibri', 10, bold=is_bold, italic=is_italic)
            
            p.paragraph_format.space_after = Pt(0 if in_references else 6)
            p.paragraph_format.line_spacing = 1.0
        
        i += 1
    
    # Save
    doc.save(output_file)
    print(f"\n✓ Converted successfully: {output_file}")
    print(f"✓ Word count: ~{len(' '.join([l.strip() for l in lines if l.strip()]).split())} words")
    print(f"\nAJET formatting applied:")
    print(f"  • Title: Arial 14pt bold, centered")
    print(f"  • Headings L1: Arial 12pt bold")  
    print(f"  • Headings L2: Arial 10pt bold")
    print(f"  • Body: Calibri 10pt")
    print(f"  • Abstract: Indented 1cm, no heading")
    print(f"  • Implications: Indented 1cm, italic heading, bullets")
    print(f"  • References: Hanging indent 0.5cm")
    print(f"  • Tables: Native Word tables with formatting")
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if skip_next:
            skip_next = False
            i += 1
            continue
        
        # Skip empty lines in references
        if in_references and not line:
            i += 1
            continue
        
        # Title (first line after ---)
        if line.startswith('# ') and not in_references:
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            set_run_font(run, 'Arial', 14, bold=True)
            set_paragraph_spacing(p, space_after=12)
            
        # Abstract section
        elif line == '## ABSTRACT':
            in_abstract = True
            # Skip the heading, AJET doesn't want "Abstract" heading
            
        elif in_abstract and line == '---':
            in_abstract = False
            
        elif in_abstract and line:
            p = doc.add_paragraph(line)
            add_left_indent(p, 1.0)
            for run in p.runs:
                set_run_font(run, 'Calibri', 10)
            set_paragraph_spacing(p, space_after=0)
            
        # Keywords
        elif line.startswith('**Keywords**:'):
            keywords_text = line.replace('**Keywords**:', '').strip()
            p = doc.add_paragraph()
            add_left_indent(p, 1.0)
            run1 = p.add_run('Keywords')
            set_run_font(run1, 'Calibri', 10, bold=True)
            run2 = p.add_run(': ' + keywords_text)
            set_run_font(run2, 'Calibri', 10)
            set_paragraph_spacing(p, space_after=12)
            
        # Implications section
        elif line == '## IMPLICATIONS FOR PRACTICE OR POLICY':
            in_implications = True
            p = doc.add_paragraph()
            add_left_indent(p, 1.0)
            run = p.add_run('Implications for practice or policy')
            set_run_font(run, 'Calibri', 10, italic=True)
            set_paragraph_spacing(p, space_after=6)
            
        elif in_implications and line == '---':
            in_implications = False
            
        elif in_implications and line.startswith('- '):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            add_left_indent(p, 1.0)
            for run in p.runs:
                set_run_font(run, 'Calibri', 10)
            set_paragraph_spacing(p, space_after=0)
            
        # References section
        elif line == '## REFERENCES':
            in_references = True
            p = doc.add_paragraph()
            run = p.add_run('REFERENCES')
            set_run_font(run, 'Arial', 12, bold=True)
            set_paragraph_spacing(p, space_after=6, space_before=12)
            
        elif in_references and line:
            # Reference entry
            p = doc.add_paragraph(line)
            add_hanging_indent(p, 0.5)
            for run in p.runs:
                set_run_font(run, 'Calibri', 10)
            set_paragraph_spacing(p, space_after=0)
            
        # Level 1 headings (##)
        elif line.startswith('## ') and not in_abstract:
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            set_run_font(run, 'Arial', 12, bold=True)
            set_paragraph_spacing(p, space_after=6, space_before=12)
            
        # Level 2 headings (###)
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            set_run_font(run, 'Arial', 10, bold=True)
            set_paragraph_spacing(p, space_after=6, space_before=6)
            
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if skip_next:
            skip_next = False
            i += 1
            continue
        
        # Skip empty lines in references
        if in_references and not line:
            i += 1
            continue
        
        # Title (first line after ---)
        if line.startswith('# ') and not in_references:
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            set_run_font(run, 'Arial', 14, bold=True)
            set_paragraph_spacing(p, space_after=12)
            
        # Abstract section
        elif line == '## ABSTRACT':
            in_abstract = True
            # Skip the heading, AJET doesn't want "Abstract" heading
            
        elif in_abstract and line == '---':
            in_abstract = False
            
        elif in_abstract and line:
            p = doc.add_paragraph(line)
            add_left_indent(p, 1.0)
            for run in p.runs:
                set_run_font(run, 'Calibri', 10)
            set_paragraph_spacing(p, space_after=0)
            
        # Keywords
        elif line.startswith('**Keywords**:'):
            keywords_text = line.replace('**Keywords**:', '').strip()
            p = doc.add_paragraph()
            add_left_indent(p, 1.0)
            run1 = p.add_run('Keywords')
            set_run_font(run1, 'Calibri', 10, bold=True)
            run2 = p.add_run(': ' + keywords_text)
            set_run_font(run2, 'Calibri', 10)
            set_paragraph_spacing(p, space_after=12)
            
        # Implications section
        elif line == '## IMPLICATIONS FOR PRACTICE OR POLICY':
            in_implications = True
            p = doc.add_paragraph()
            add_left_indent(p, 1.0)
            run = p.add_run('Implications for practice or policy')
            set_run_font(run, 'Calibri', 10, italic=True)
            set_paragraph_spacing(p, space_after=6)
            
        elif in_implications and line == '---':
            in_implications = False
            
        elif in_implications and line.startswith('- '):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            add_left_indent(p, 1.0)
            for run in p.runs:
                set_run_font(run, 'Calibri', 10)
            set_paragraph_spacing(p, space_after=0)
            
        # References section
        elif line == '## REFERENCES':
            in_references = True
            p = doc.add_paragraph()
            run = p.add_run('REFERENCES')
            set_run_font(run, 'Arial', 12, bold=True)
            set_paragraph_spacing(p, space_after=6, space_before=12)
            
        elif in_references and line:
            # Reference entry
            p = doc.add_paragraph(line)
            add_hanging_indent(p, 0.5)
            for run in p.runs:
                set_run_font(run, 'Calibri', 10)
            set_paragraph_spacing(p, space_after=0)
            
        # Level 1 headings (##)
        elif line.startswith('## ') and not in_abstract:
            heading_text = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            set_run_font(run, 'Arial', 12, bold=True)
            set_paragraph_spacing(p, space_after=6, space_before=12)
            
        # Level 2 headings (###)
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            set_run_font(run, 'Arial', 10, bold=True)
            set_paragraph_spacing(p, space_after=6, space_before=6)
        
        # Markdown tables (check for table start)
        elif line.startswith('|') and not in_references:
            table_data, end_idx = parse_markdown_table(lines, i)
            if table_data:
                header, rows = table_data
                add_word_table(doc, header, rows)
                i = end_idx
                continue
        
        # Skip horizontal rules
        elif line == '---':
            pass
            
        # Regular paragraphs
        elif line and not line.startswith('#'):
            p = doc.add_paragraph()
            
            # Process markdown formatting
            segments = process_markdown_formatting(line)
            for text, is_bold, is_italic in segments:
                run = p.add_run(text)
                set_run_font(run, 'Calibri', 10, bold=is_bold, italic=is_italic)
            
            set_paragraph_spacing(p, space_after=6)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"\n✓ Document saved: {output_file}")
    print(f"✓ AJET formatting applied:")
    print(f"  - Title: Arial 14pt bold, centered")
    print(f"  - Headings L1: Arial 12pt bold")
    print(f"  - Headings L2: Arial 10pt bold")
    print(f"  - Body: Calibri 10pt, left-justified")
    print(f"  - Margins: 3.0 cm all sides")
    print(f"  - Abstract: Indented 1.0 cm, no heading")
    print(f"  - Implications: Indented 1.0 cm, italic heading")
    print(f"  - References: Hanging indent 0.5 cm")

if __name__ == '__main__':
    import os
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, 'MANUSCRIPT_DRAFT_v1.md')
    output_file = os.path.join(script_dir, 'MANUSCRIPT_AJET_FORMATTED.docx')
    
    print("\n" + "="*70)
    print("  CONVERTING TO AJET-FORMATTED DOCX")
    print("="*70)
    print(f"\nInput: {os.path.basename(md_file)}")
    print(f"Output: {os.path.basename(output_file)}")
    
    convert_markdown_to_docx(md_file, output_file)
    
    print("\n" + "="*70)
    print("  CONVERSION COMPLETE")
    print("="*70)
    print("\nNext steps:")
    print("  1. Open MANUSCRIPT_AJET_FORMATTED.docx")
    print("  2. Insert 8 figures from osf_upload/03_Results/figures/")
    print("  3. Review formatting and adjust if needed")
    print("  4. Strip metadata before submission")
    print("\n")
