"""
Simple and robust Markdown to AJET-formatted Word converter
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_run_font(run, font_name, font_size, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # Ensure font applies to all character types
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def add_formatted_text(paragraph, text):
    """Add text with markdown bold/italic formatting"""
    # Split by ** for bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            set_run_font(run, 'Calibri', 10, bold=True)
        elif part:
            # Check for italic
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    run = paragraph.add_run(ipart[1:-1])
                    set_run_font(run, 'Calibri', 10, italic=True)
                elif ipart:
                    run = paragraph.add_run(ipart)
                    set_run_font(run, 'Calibri', 10)

def parse_table(lines, start_idx):
    """Parse markdown table"""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1
    
    if len(table_lines) < 3:
        return None, start_idx
    
    # Header
    header = [c.strip() for c in table_lines[0].split('|')[1:-1]]
    # Rows (skip separator at index 1)
    rows = []
    for line in table_lines[2:]:
        row = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(row)
    
    return (header, rows), i

def add_table(doc, header, rows):
    """Add formatted table"""
    table = doc.add_table(rows=1+len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    
    # Header
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, 'Calibri', 10, bold=True)
    
    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, 'Calibri', 10)

def convert_to_docx(md_file, docx_file):
    """Main conversion function"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = [line.rstrip() for line in f.readlines()]
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
    
    # State
    in_abstract = False
    in_implications = False
    in_references = False
    skip_research_highlights = False
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty
        if not line:
            i += 1
            continue
        
        # Skip horizontal rules
        if line == '---':
            i += 1
            continue
        
        # Title (first # heading)
        if line.startswith('# ') and i < 10:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            set_run_font(run, 'Arial', 14, bold=True)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        
        # Skip Running Title
        if 'Running Title' in line:
            i += 1
            continue
        
        # ABSTRACT
        if line == '## ABSTRACT':
            in_abstract = True
            # No heading for abstract per AJET
            i += 1
            continue
        
        # IMPLICATIONS
        if line == '## IMPLICATIONS FOR PRACTICE OR POLICY':
            in_abstract = False
            in_implications = True
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            run = p.add_run('Implications for practice or policy')
            set_run_font(run, 'Calibri', 10, italic=True)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # RESEARCH HIGHLIGHTS - skip
        if line == '## RESEARCH HIGHLIGHTS':
            skip_research_highlights = True
            i += 1
            continue
        
        # End of research highlights
        if skip_research_highlights and line.startswith('## '):
            skip_research_highlights = False
            # Will be processed below
        
        if skip_research_highlights:
            i += 1
            continue
        
        # REFERENCES
        if line == '## REFERENCES':
            in_implications = False
            in_references = True
            p = doc.add_paragraph()
            run = p.add_run('REFERENCES')
            set_run_font(run, 'Arial', 12, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Level 1 heading (##)
        if line.startswith('## '):
            in_abstract = False
            in_implications = False
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            set_run_font(run, 'Arial', 12, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Level 2 heading (###)
        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            set_run_font(run, 'Arial', 10, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Table
        if line.startswith('|'):
            table_data, end_idx = parse_table(lines, i)
            if table_data:
                header, rows = table_data
                add_table(doc, header, rows)
                i = end_idx
                continue
        
        # Figure placeholders - add text placeholder
        if line.startswith('!['):
            # Extract figure caption and filename
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                caption_text = match.group(1)
                filename = match.group(2).split('/')[-1]
                
                # Add placeholder paragraph
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[INSERT FIGURE HERE: {filename}]")
                set_run_font(run, 'Calibri', 10, italic=True)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.space_before = Pt(6)
                
                # Add caption
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(caption_text)
                set_run_font(run, 'Calibri', 10, italic=True)
                p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        
        # Bullet points (implications)
        if in_implications and line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            add_formatted_text(p, line[2:])
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        
        # Skip research highlights bullets
        if line.startswith('•'):
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        
        # Indent abstract
        if in_abstract:
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
        
        # Hanging indent references
        if in_references:
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
        
        # Add text
        add_formatted_text(p, line)
        
        # Spacing
        p.paragraph_format.space_after = Pt(0 if in_references else 6)
        
        i += 1
    
    doc.save(docx_file)
    print(f"\n✓ Converted: {docx_file}")
    print(f"\nAJET formatting applied:")
    print(f"  • Title: Arial 14pt bold, centered")
    print(f"  • Headings L1: Arial 12pt bold")
    print(f"  • Headings L2: Arial 10pt bold")
    print(f"  • Body: Calibri 10pt")
    print(f"  • Abstract: Indented 1cm")
    print(f"  • Implications: Indented 1cm, italic heading")
    print(f"  • References: Hanging indent 0.5cm")
    print(f"  • Tables: Native Word tables\n")

if __name__ == '__main__':
    import os
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, 'MANUSCRIPT_DRAFT_v1.md')
    docx_file = os.path.join(script_dir, 'MANUSCRIPT_AJET_FORMATTED.docx')
    
    print("\n" + "="*70)
    print("  CONVERTING MANUSCRIPT TO AJET WORD FORMAT")
    print("="*70)
    
    convert_to_docx(md_file, docx_file)
    
    print("="*70)
    print("  NEXT STEPS:")
    print("="*70)
    print("  1. Open MANUSCRIPT_AJET_FORMATTED.docx")
    print("  2. Insert 8 figures from ../osf_upload/03_Results/figures/")
    print("  3. Review all formatting")
    print("  4. Strip metadata (File → Inspect Document → Remove All)")
    print("="*70 + "\n")
